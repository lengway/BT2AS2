from docx import Document
from pathlib import Path

p = Path("c:/Users/lengway/Desktop/Projects/BT2/AS2/docs/Assignment 1 Roman Kutbeyev.docx")
d = Document(p)
print("PARAGRAPHS:", len(d.paragraphs))
for i, para in enumerate(d.paragraphs, 1):
    text = para.text.strip()
    if text:
        style = para.style.name if para.style else "N/A"
        print(f"{i:03d} [{style}] {text}")
