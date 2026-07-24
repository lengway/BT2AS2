from docx import Document
from docx.shared import Pt
from datetime import date
from pathlib import Path


def add_task_block(doc: Document, title: str, done: str, why: str, result: str, screenshot: str) -> None:
    doc.add_heading(title, level=2)
    doc.add_paragraph("Что сделано:", style="List Bullet")
    doc.add_paragraph(done, style="List Bullet 2")
    doc.add_paragraph("Зачем:", style="List Bullet")
    doc.add_paragraph(why, style="List Bullet 2")
    doc.add_paragraph("Что это дало:", style="List Bullet")
    doc.add_paragraph(result, style="List Bullet 2")
    doc.add_paragraph("Какой скрин вставить:", style="List Bullet")
    doc.add_paragraph(screenshot, style="List Bullet 2")


root = Path("c:/Users/lengway/Desktop/Projects/BT2/AS2")
out = root / "docs" / "Assignment_2_Short_Report_v2.docx"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

doc.add_heading("Assignment 2 — Краткий отчет по выполнению", 0)
doc.add_paragraph(f"Дата: {date.today().isoformat()}")
doc.add_paragraph("Студент: ____________________")

doc.add_heading("Структура проекта", level=1)
for p in [
    "part1/task1",
    "part1/task2",
    "part2/task3",
    "part2/task4",
    "part3/task5",
    "part4/task6",
]:
    doc.add_paragraph(p, style="List Bullet")

doc.add_heading("Кратко по задачам", level=1)

add_task_block(
    doc,
    "Part 1 — Task 1 (ERC-20 + unit/fuzz/invariant)",
    "Собран отдельный Foundry-проект, реализован ERC-20, добавлены unit-тесты, fuzz-тест transfer, 2 invariant-теста и прогон coverage.",
    "Чтобы проверить и точечную корректность логики, и устойчивость на случайных входах.",
    "Покрыты базовые функции токена и крайние случаи; получен coverage-отчет для подтверждения качества тестирования.",
    "Скрин терминала с PASS по test + кусок таблицы coverage из part1/task1/outputs/coverage.txt.",
)

add_task_block(
    doc,
    "Part 1 — Task 2 (Fork testing)",
    "Добавлены fork-тесты: чтение totalSupply USDC и симуляция swap через Uniswap V2 router; добавлено пояснение по vm.createSelectFork/vm.rollFork.",
    "Чтобы показать тестирование на реальных контрактах и реальном состоянии сети, а не только на моках.",
    "Есть готовая заготовка для интеграционных fork-проверок; при наличии MAINNET_RPC_URL тесты запускаются на реальном форке.",
    "Скрин запуска part1/task2/test/ForkMainnet.t.sol (лучше с реальным MAINNET_RPC_URL), где видно PASS по fork-тестам.",
)

add_task_block(
    doc,
    "Part 2 — Task 3 (AMM)",
    "Реализован AMM по x*y=k: addLiquidity, removeLiquidity, swap с fee 0.3%, slippage protection, LP token, события; написан полный набор тестов (17).",
    "Чтобы реализовать рабочее ядро DEX/AMM и проверить его поведение в практических сценариях.",
    "Получен рабочий протокол с проверенной логикой ликвидности/свопов, включая edge cases, fuzz и проверку инварианта k.",
    "Скрин PASS по AMM тестам + скрин/фрагмент gas report из part2/task3/outputs/gas-report.txt.",
)

add_task_block(
    doc,
    "Part 2 — Task 4 (Math analysis)",
    "Подготовлен технический документ с формулами: вывод constant product, влияние комиссии, IL для 2x изменения цены, price impact, сравнение с Uniswap V2.",
    "Чтобы теоретически обосновать модель AMM и показать понимание экономики пула.",
    "Есть законченная аналитическая часть, напрямую связанная с реализованным контрактом AMM.",
    "Скрин первой страницы/главных формул из part2/task4/AMM-mathematical-analysis.md.",
)

add_task_block(
    doc,
    "Part 3 — Task 5 (Lending pool)",
    "Реализован lending pool: deposit/borrow/repay/withdraw/liquidate, LTV 75%, health factor, линейные проценты, тесты с vm.warp; добавлена workflow-диаграмма.",
    "Чтобы показать базовую механику протокола кредитования и управление риском позиции.",
    "Проверен полный жизненный цикл позиции пользователя, включая ликвидацию при падении цены и начисление процентов.",
    "Скрин PASS по part3/task5 тестам + фрагмент gas report, отдельно можно скрин mermaid-диаграммы.",
)

add_task_block(
    doc,
    "Part 4 — Task 6 (CI/CD)",
    "Создан workflow test.yml: установка Foundry, build/test/coverage/gas-report, шаг статанализа Slither; добавлено краткое описание стадий.",
    "Чтобы автоматизировать проверку контрактов и ловить проблемы до деплоя.",
    "Получен базовый CI-шаблон для continuous validation смарт-контрактов.",
    "Скрин успешного GitHub Actions job (или локального прогона с аналогичным логом стадий).",
)

doc.add_heading("Итог", level=1)
doc.add_paragraph("Деплой не выполнялся (по требованию). Во всех релевантных задачах подготовлен фундамент под Sepolia: deploy-скрипты и .env.example с нужными переменными.")

out.parent.mkdir(parents=True, exist_ok=True)
doc.save(out)
print(f"Created: {out}")
