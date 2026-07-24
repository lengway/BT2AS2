from docx import Document
from pathlib import Path

out = Path("c:/Users/lengway/Desktop/Projects/BT2/AS2/docs/Assignment 2 Roman Kutbeyev.docx")

doc = Document()

def n(text: str):
    doc.add_paragraph(text)

def b(text: str):
    doc.add_paragraph(text, style="List Paragraph")


def shot(label: str, where: str):
    b(f"[INSERT SCREENSHOT: {label}]")
    b(f"Source: {where}")

n("Roman Kutbeyev")
n("SE-2404")
n("Blockchain Technologies 2")
n("Assignment 2")

n("Part 1.")
n("Task 1.")
b("Implemented separate Foundry project in part1/task1 with src/, test/, script/ structure.")
b("Implemented SimpleERC20.sol with mint, transfer, approve, transferFrom and custom errors for edge cases.")
b("Implemented unit tests covering mint, transfer, approve, transferFrom and revert scenarios (13 unit tests).")
b("Implemented fuzz test for transfer and invariant tests for supply/balance consistency.")
b("Generated coverage report and saved outputs in part1/task1/outputs.")
b("Here we can see passing test output and coverage table in: part1/task1/outputs/test.txt and part1/task1/outputs/coverage.txt")
shot("Task 1 tests passed", "Terminal run in part1/task1 with forge test -vv")
shot("Task 1 coverage table", "part1/task1/outputs/coverage.txt")

n("Task 2.")
b("Implemented fork testing project in part1/task2 with test/ForkMainnet.t.sol.")
b("Added test that reads real USDC totalSupply from mainnet USDC contract.")
b("Added test that simulates Uniswap V2 swap using real router interface on fork.")
b("Added explanation document for vm.createSelectFork and vm.rollFork behavior.")
b("Fork tests are environment-ready: with MAINNET_RPC_URL set, tests run against real chain state.")
b("Here we can see fork test script and output in: part1/task2/test/ForkMainnet.t.sol and part1/task2/outputs/test.txt")
shot("Task 2 fork test output", "Terminal run in part1/task2 with MAINNET_RPC_URL configured")
shot("USDC totalSupply + Uniswap swap test lines", "part1/task2/outputs/test.txt")

n("Part 2.")
n("Task 3.")
b("Implemented AMM.sol based on constant product x * y = k.")
b("Implemented LPToken.sol and two ERC-20 pair tokens (TokenA, TokenB).")
b("Implemented addLiquidity, removeLiquidity, swap with 0.3% fee and slippage protection.")
b("Implemented events: LiquidityAdded, LiquidityRemoved, Swap.")
b("Implemented comprehensive test suite including first/subsequent liquidity, both swap directions, slippage reverts, edge cases, large trade impact and fuzz swap test.")
b("Result: 17/17 tests passing and gas report generated.")
b("Here we can see test and gas results in: part2/task3/outputs/test.txt and part2/task3/outputs/gas-report.txt")
shot("Task 3 tests passed", "part2/task3/outputs/test.txt")
shot("Task 3 gas report table", "part2/task3/outputs/gas-report.txt")

n("Task 4.")
b("Prepared technical mathematical analysis in part2/task4/AMM-mathematical-analysis.md.")
b("Included derivation of constant product formula and pricing logic.")
b("Included explanation how 0.3% fee affects invariant k over time.")
b("Included impermanent loss derivation and IL calculation for 2x price change.")
b("Included price impact analysis by trade size and comparison with Uniswap V2 missing features.")
b("Here we can see the analysis document in: part2/task4/AMM-mathematical-analysis.md")
shot("Task 4 analysis preview", "Open part2/task4/AMM-mathematical-analysis.md and capture first formula section")

n("Part 3.")
n("Task 5.")
b("Implemented LendingPool.sol with deposit, borrow, repay, withdraw and liquidate.")
b("Implemented 75% LTV rule, health factor checks and position tracking (deposited/borrowed/lastAccrued).")
b("Implemented simple linear interest accrual model and tested time-based accrual with vm.warp.")
b("Implemented liquidation scenario with oracle price drop simulation.")
b("Implemented complete test suite (11 tests) and generated gas report.")
b("Added workflow diagram in part3/task5/docs/workflow-diagram.md.")
b("Here we can see lending tests and gas in: part3/task5/outputs/test.txt and part3/task5/outputs/gas-report.txt")
shot("Task 5 tests passed", "part3/task5/outputs/test.txt")
shot("Task 5 gas report table", "part3/task5/outputs/gas-report.txt")
shot("Task 5 workflow diagram", "part3/task5/docs/workflow-diagram.md")

n("Part 4.")
n("Task 6.")
b("Implemented CI pipeline YAML in part4/task6/.github/workflows/test.yml.")
b("Pipeline includes: Foundry install, build, test, coverage, gas report and Slither analysis steps.")
b("Added short documentation of stages in part4/task6/pipeline-stages.md.")
b("This gives automated quality/security checks before deployment.")
b("Here we can see pipeline files in: part4/task6/.github/workflows/test.yml and part4/task6/pipeline-stages.md")
shot("Task 6 CI workflow YAML", "part4/task6/.github/workflows/test.yml")
shot("Task 6 pipeline stages doc", "part4/task6/pipeline-stages.md")

b("No deployment was performed, as required. Sepolia deployment foundation (scripts + .env.example) was prepared in relevant tasks.")

out.parent.mkdir(parents=True, exist_ok=True)
doc.save(out)
print(f"Created: {out}")
